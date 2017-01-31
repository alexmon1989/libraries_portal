from openpyxl import Workbook
from openpyxl.styles import Font, Alignment

from docx import Document

from wkhtmltopdf.views import PDFTemplateResponse

from pyexcel_ods import save_data

from orders.models import OrderAction

from abc import ABCMeta, abstractmethod
from datetime import datetime
import os


class OrdersReportStrategy:
    """Абстрактный класс стратегии формирования отчёта."""
    __metaclass__ = ABCMeta

    @abstractmethod
    def make(self, data):
        pass


class OrdersReportXLSStrategy(OrdersReportStrategy):
    """Стратегия формирования отчёта в формате XLSX."""
    def make(self, data):
        filename = 'orders_acquisitions_{}.xlsx'.format(datetime.today().timestamp())
        filepath = os.path.abspath(
            os.path.join('static', 'uploads', 'reports', filename)
        )

        wb = Workbook()
        ws = wb.active

        f = Font(bold=True)
        al = Alignment(horizontal="center", vertical="center")

        ws['A1'].font = f
        ws['A1'].alignment = al
        ws['A1'] = 'Номер заказа'

        ws['B1'] = 'Читатель'
        ws['B1'].font = f
        ws['B1'].alignment = al

        ws['C1'] = 'Заказ на'
        ws['C1'].font = f
        ws['C1'].alignment = al

        ws['D1'] = 'Действие'
        ws['D1'].font = f
        ws['D1'].alignment = al

        ws['E1'] = 'Дата'
        ws['E1'].font = f
        ws['E1'].alignment = al

        ws.column_dimensions['A'].width = 50
        ws.column_dimensions['B'].width = 50
        ws.column_dimensions['C'].width = 50
        ws.column_dimensions['D'].width = 50
        ws.column_dimensions['E'].width = 50

        for idx, item in enumerate(data):
            ws['A{}'.format(idx+2)] = str(item[0])
            ws['B{}'.format(idx+2)] = str(item[1])
            ws['C{}'.format(idx+2)] = str(item[2])
            ws['D{}'.format(idx+2)] = str(item[3])
            ws['E{}'.format(idx+2)] = item[4].strftime("%d.%m.%Y %H:%M:%S")

        wb.save(filename=filepath)

        return filepath


class OrdersReportDOCStrategy(OrdersReportStrategy):
    """Стратегия формирования отчёта в формате DOCX."""
    def make(self, data):
        filename = 'orders_acquisitions_{}.docx'.format(datetime.today().timestamp())
        filepath = os.path.abspath(
            os.path.join('static', 'uploads', 'reports', filename)
        )

        document = Document()

        document.add_heading('Отчёт "Обработанные заказы"')

        col_names = ('Номер заказа',
                     'Читатель',
                     'Заказ на',
                     'Действие',
                     'Дата')
        table = document.add_table(rows=1, cols=len(col_names), style="TableGrid")
        hdr_cells = table.rows[0].cells
        for idx, name in enumerate(col_names):
            paragraph = hdr_cells[idx].paragraphs[0]
            run = paragraph.add_run(name)
            run.bold = True

        for item in data:
            cells = table.add_row().cells
            cells[0].text = str(item[0])
            cells[1].text = str(item[1])
            cells[2].text = str(item[2])
            cells[3].text = str(item[3])
            cells[4].text = item[4].strftime("%d.%m.%Y %H:%M:%S")

        document.save(filepath)

        return filepath


class OrdersReportPDFStrategy(OrdersReportStrategy):
    """Стратегия формирования отчёта в формате PDF."""
    def make(self, data):
        filename = 'orders_acquisitions_{}.pdf'.format(datetime.today().timestamp())
        filepath = os.path.abspath(
            os.path.join('static', 'uploads', 'reports', filename)
        )

        response = PDFTemplateResponse(
            request=None,
            template='library_cabinet/reports/pdf_templates/orders_acquisitions.html',
            context={'items': data},
            cmd_options={'load-error-handling': 'ignore'})

        with open(filepath, "wb") as f:
            f.write(response.rendered_content)

        return filepath


class OrdersReportODSStrategy(OrdersReportStrategy):
    """Стратегия формирования отчёта в формате ODS."""
    def make(self, data):
        filename = 'orders_acquisitions_{}.ods'.format(datetime.today().timestamp())
        filepath = os.path.abspath(
            os.path.join('static', 'uploads', 'reports', filename)
        )

        # Формирование данных
        table_data = [['Номер заказа',
                       'Читатель',
                       'Заказ на',
                       'Дата заявки',
                       'Статус']]
        for item in data:
            table_data.append([
                str(item[0]),
                str(item[1]),
                str(item[2]),
                item[3].strftime("%d.%m.%Y %H:%M:%S"),
                str(item[4]),
            ])

        data = {"Sheet 1": table_data}
        save_data(filepath, data)

        return filepath


class OrdersReportContext:
    """Контекст использования стратегии."""
    strategy = None
    params = {}

    def __init__(self, strategy, params, user):
        """Конструктор."""
        self.strategy = strategy
        self.params = params
        self.user = user

    def get_data(self):
        """Получает данные, из которых формируется отчёт."""
        data = []
        orders_actions = OrderAction.objects.filter(
            user=self.user,
            order__user_id__in=self.params.getlist('reader'),
            action_date__gte=datetime.strptime(self.params['date_from'], '%d.%m.%Y'),
            action_date__lte=datetime.strptime(self.params['date_to'], '%d.%m.%Y'),
            status_id__in=self.params.getlist('status')
        ).order_by('-action_date')

        # Формирование данных для отчёта
        for order_action in orders_actions:
            data.append([
                order_action.order.pk,
                '{} {} (г. {}, читательский билет № {})'.format(order_action.order.user.last_name,
                                                                order_action.order.user.first_name,
                                                                order_action.order.user.profile.city,
                                                                order_action.order.user.readerticket.id),
                order_action.order.document,
                order_action.status,
                order_action.action_date
            ])

        return data

    def get_report(self):
        """Формирует отчёт."""
        return self.strategy.make(self.get_data())
