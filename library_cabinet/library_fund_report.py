from django.contrib.auth.models import User

from openpyxl import Workbook
from openpyxl.styles import Font, Alignment

from docx import Document

from pyexcel_ods import save_data

from wkhtmltopdf.views import PDFTemplateResponse

from orders.models import OrderAction
from home.models import LibraryRequestAccess

from abc import ABCMeta, abstractmethod
from datetime import datetime
import os


class LibraryFundReportStrategy:
    """Абстрактный класс стратегии формирования отчёта."""
    __metaclass__ = ABCMeta

    @abstractmethod
    def make(self, data):
        pass


class LibraryFundReportXLSStrategy(LibraryFundReportStrategy):
    """Стратегия формирования отчёта в формате XLSX."""
    def make(self, data):
        filename = 'library_fund_{}.xlsx'.format(datetime.today().timestamp())
        filepath = os.path.abspath(
            os.path.join('static', 'uploads', 'reports', filename)
        )

        wb = Workbook()
        ws = wb.active

        f = Font(bold=True)
        al = Alignment(horizontal="center", vertical="center")

        ws['A1'].font = f
        ws['A1'].alignment = al
        ws['A1'] = 'Сотрудник'

        ws['B1'] = 'Количество выданных книг'
        ws['B1'].font = f
        ws['B1'].alignment = al

        ws['C1'] = 'Количество зарегистрированных читателей'
        ws['C1'].font = f
        ws['C1'].alignment = al

        ws['D1'] = 'Зарегистрированные читатели'
        ws['D1'].font = f
        ws['D1'].alignment = al

        ws.column_dimensions['A'].width = 50
        ws.column_dimensions['B'].width = 50
        ws.column_dimensions['C'].width = 50
        ws.column_dimensions['D'].width = 50

        for idx, item in enumerate(data):
            ws['A{}'.format(idx+2)] = item[0]
            ws['B{}'.format(idx+2)] = item[1]
            ws['C{}'.format(idx+2)] = item[2]
            ws['D{}'.format(idx+2)].alignment = Alignment(wrapText=True)
            ws['D{}'.format(idx+2)] = item[3]

        wb.save(filename=filepath)

        return filepath


class LibraryFundReportDOCStrategy(LibraryFundReportStrategy):
    """Стратегия формирования отчёта в формате DOCX."""
    def make(self, data):
        filename = 'library_fund_{}.docx'.format(datetime.today().timestamp())
        filepath = os.path.abspath(
            os.path.join('static', 'uploads', 'reports', filename)
        )

        document = Document()

        document.add_heading('Отчёт "Библиотечный фонд"')

        col_names = ('Сотрудник',
                     'Количество выданных книг',
                     'Количество зарегистрированных читателей',
                     'Зарегистрированные читатели')
        table = document.add_table(rows=1, cols=len(col_names), style="TableGrid")
        hdr_cells = table.rows[0].cells
        for idx, name in enumerate(col_names):
            paragraph = hdr_cells[idx].paragraphs[0]
            run = paragraph.add_run(name)
            run.bold = True

        for item in data:
            cells = table.add_row().cells
            cells[0].text = item[0]
            cells[1].text = str(item[1])
            cells[2].text = str(item[2])
            cells[3].text = item[3]

        document.save(filepath)

        return filepath


class LibraryFundReportPDFStrategy(LibraryFundReportStrategy):
    """Стратегия формирования отчёта в формате PDF."""
    def make(self, data):
        filename = 'library_fund_{}.pdf'.format(datetime.today().timestamp())
        filepath = os.path.abspath(
            os.path.join('static', 'uploads', 'reports', filename)
        )

        response = PDFTemplateResponse(
            request=None,
            template='library_cabinet/reports/pdf_templates/library_fund.html',
            context={'items': data},
            cmd_options={'load-error-handling': 'ignore'})

        with open(filepath, "wb") as f:
            f.write(response.rendered_content)

        return filepath


class LibraryFundReportODSStrategy(LibraryFundReportStrategy):
    """Стратегия формирования отчёта в формате ODS."""
    def make(self, data):
        filename = 'library_fund_{}.ods'.format(datetime.today().timestamp())
        filepath = os.path.abspath(
            os.path.join('static', 'uploads', 'reports', filename)
        )

        # Формирование данных
        table_data = [['Сотрудник',
                       'Количество выданных книг',
                       'Количество зарегистрированных читателей',
                       'Зарегистрированные читатели']]
        for item in data:
            table_data.append([
                item[0],
                item[1],
                item[2],
                item[3]
            ])

        data = {"Sheet 1": table_data}
        save_data(filepath, data)

        return filepath


class LibraryFundReportContext:
    """Контекст использования стратегии."""
    strategy = None
    params = {}

    def __init__(self, strategy, params):
        """Конструктор."""
        self.strategy = strategy
        self.params = params

    def get_data(self):
        """Получает данные, из которых формируется отчёт."""
        data = []
        for user_id in self.params.getlist('user'):
            user = User.objects.get(id=user_id)
            # Имя сотрудника
            user_data = ['{} {}'.format(user.last_name, user.first_name)]

            # Выданные заказы
            orders = OrderAction.objects.filter(
                user_id=user_id,
                status__title='Выдан',
                action_date__gte=datetime.strptime(self.params['date_from'], '%d.%m.%Y'),
                action_date__lte=datetime.strptime(self.params['date_to'], '%d.%m.%Y')
            ).order_by('-action_date')

            user_data.append(orders.count())

            # Зарегистрированные читатели (одобренные заявки)
            approved_users = LibraryRequestAccess.objects.filter(approved_by=user_id,
                                                                 updated_at__gte=datetime.strptime(
                                                                     self.params['date_from'],
                                                                     '%d.%m.%Y'),
                                                                 updated_at__lte=datetime.strptime(
                                                                     self.params['date_to'],
                                                                     '%d.%m.%Y'))
            user_data.append(approved_users.count())
            approved_users_list = []
            for approved_user in approved_users:
                approved_users_list\
                    .append('{} {} (г. {}, читательский билет № {})'
                            .format(approved_user.user.last_name,
                                    approved_user.user.first_name,
                                    approved_user.user.profile.city,
                                    approved_user.user.readerticket.id))
            user_data.append(", ".join(approved_users_list))

            data.append(user_data)

        return data

    def get_report(self):
        """Формирует отчёт."""
        return self.strategy.make(self.get_data())
