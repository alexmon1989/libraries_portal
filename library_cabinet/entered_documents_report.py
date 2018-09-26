from django.conf import settings

from openpyxl import Workbook
from openpyxl.styles import Font, Alignment

from docx import Document

from wkhtmltopdf.views import PDFTemplateResponse

from pyexcel_ods import save_data

from documents.models import Document

from abc import ABCMeta, abstractmethod
from datetime import datetime
import os


class EnteredDocumentsReportStrategy:
    """Абстрактный класс стратегии формирования отчёта."""
    __metaclass__ = ABCMeta

    @abstractmethod
    def make(self, data):
        pass


class EnteredDocumentsReportXLSStrategy(EnteredDocumentsReportStrategy):
    """Стратегия формирования отчёта в формате XLSX."""
    def make(self, data):
        filename = 'entered_documents_{}.xlsx'.format(datetime.today().timestamp())
        filepath = os.path.abspath(
            os.path.join('static', 'uploads', 'reports', filename)
        )

        wb = Workbook()
        ws = wb.active

        f = Font(bold=True)
        al = Alignment(horizontal="center", vertical="center")

        ws['A1'].font = f
        ws['A1'].alignment = al
        ws['A1'] = 'Документ'

        ws['B1'] = 'Дата внесения'
        ws['B1'].font = f
        ws['B1'].alignment = al

        ws.column_dimensions['A'].width = 50
        ws.column_dimensions['B'].width = 50

        for idx, item in enumerate(data):
            ws['A{}'.format(idx+2)] = str(item[0])
            ws['B{}'.format(idx+2)] = str(item[1])

        wb.save(filename=filepath)

        return filepath


class EnteredDocumentsReportDOCStrategy(EnteredDocumentsReportStrategy):
    """Стратегия формирования отчёта в формате DOCX."""
    def make(self, data):
        filename = 'entered_documents_{}.docx'.format(datetime.today().timestamp())
        filepath = os.path.join(settings.MEDIA_ROOT, 'reports', filename)

        document = Document()

        document.add_heading('Отчёт "Обработанные заказы"')

        col_names = ('Документ',
                     'Дата внесения')
        table = document.add_table(rows=1, cols=len(col_names), style="TableGrid")
        hdr_cells = table.rows[0].cells
        for idx, name in enumerate(col_names):
            paragraph = hdr_cells[idx].paragraphs[0]
            run = paragraph.add_run(name)
            run.bold = True

        for item in data:
            cells = table.add_row().cells
            cells[0].text = str(item[0])
            cells[1].text = str(item[1])

        document.save(filepath)

        return filepath


class EnteredDocumentsReportPDFStrategy(EnteredDocumentsReportStrategy):
    """Стратегия формирования отчёта в формате PDF."""
    def make(self, data):
        filename = 'entered_documents_{}.pdf'.format(datetime.today().timestamp())
        filepath = os.path.abspath(
            os.path.join('static', 'uploads', 'reports', filename)
        )

        response = PDFTemplateResponse(
            request=None,
            template='library_cabinet/reports/pdf_templates/entered_documents.html',
            context={'items': data},
            cmd_options={'load-error-handling': 'ignore'})

        with open(filepath, "wb") as f:
            f.write(response.rendered_content)

        return filepath


class EnteredDocumentsReportODSStrategy(EnteredDocumentsReportStrategy):
    """Стратегия формирования отчёта в формате ODS."""
    def make(self, data):
        filename = 'entered_documents_{}.ods'.format(datetime.today().timestamp())
        filepath = os.path.abspath(
            os.path.join('static', 'uploads', 'reports', filename)
        )

        # Формирование данных
        table_data = [['Документ',
                       'Дата внесения']]
        for item in data:
            table_data.append([
                str(item[0]),
                str(item[1])
            ])

        data = {"Sheet 1": table_data}
        save_data(filepath, data)

        return filepath


class EnteredDocumentsReportContext:
    """Контекст использования стратегии."""
    strategy = None
    params = {}

    def __init__(self, strategy, params, user):
        """Конструктор."""
        self.strategy = strategy
        self.params = params
        self.user = user

    def get_data(self):
        """Получает данные, из которых формируется отчёт."""
        data = []
        documents = Document.objects.filter(
            created_by_id=self.user,
            created_at__gte=datetime.strptime(self.params['date_from'], '%d.%m.%Y'),
            created_at__lte=datetime.strptime(self.params['date_to'], '%d.%m.%Y')
        ).order_by('-created_at')

        # Формирование данных для отчёта
        for document in documents:
            data.append([
                document,
                document.created_at.strftime("%d.%m.%Y %H:%M:%S")
            ])

        return data

    def get_report(self):
        """Формирует отчёт."""
        return self.strategy.make(self.get_data())
