from django.contrib.auth.models import User

from openpyxl import Workbook
from openpyxl.styles import Font, Alignment

from docx import Document

from pyexcel_ods import save_data

from wkhtmltopdf.views import PDFTemplateResponse

from orders.models import OrderAction
from documents.models import Document as DocumentModel

from abc import ABCMeta, abstractmethod
from datetime import datetime
import os


class StaffReportStrategy:
    """Абстрактный класс стратегии формирования отчёта."""
    __metaclass__ = ABCMeta

    @abstractmethod
    def make(self, data):
        pass


class StaffReportXLSStrategy(StaffReportStrategy):
    """Стратегия формирования отчёта в формате XLSX."""
    def make(self, data):
        filename = 'staff_{}.xlsx'.format(datetime.today().timestamp())
        filepath = os.path.abspath(
            os.path.join('static', 'uploads', 'reports', filename)
        )

        wb = Workbook()
        ws = wb.active

        f = Font(bold=True)
        al = Alignment(horizontal="center", vertical="center")

        ws['A1'].font = f
        ws['A1'].alignment = al
        ws['A1'] = 'Сотрудник'

        ws['B1'] = 'Занесено книг в базу данных'
        ws['B1'].font = f
        ws['B1'].alignment = al

        ws['C1'] = 'Обслужено читателей'
        ws['C1'].font = f
        ws['C1'].alignment = al

        ws['D1'] = 'Занесенные книги в базу данных'
        ws['D1'].font = f
        ws['D1'].alignment = al

        ws['E1'] = 'Обслуженные читатели'
        ws['E1'].font = f
        ws['E1'].alignment = al

        ws.column_dimensions['A'].width = 50
        ws.column_dimensions['B'].width = 50
        ws.column_dimensions['C'].width = 50
        ws.column_dimensions['D'].width = 50
        ws.column_dimensions['E'].width = 50

        for idx, item in enumerate(data):
            ws['A{}'.format(idx+2)] = str(item[0])
            ws['B{}'.format(idx+2)] = str(item[1])
            ws['C{}'.format(idx+2)] = str(item[2])
            ws['D{}'.format(idx+2)] = str(item[3])
            ws['D{}'.format(idx+2)].alignment = Alignment(wrapText=True)
            ws['E{}'.format(idx+2)] = str(item[4])
            ws['E{}'.format(idx+2)].alignment = Alignment(wrapText=True)

        wb.save(filename=filepath)

        return filepath


class StaffReportDOCStrategy(StaffReportStrategy):
    """Стратегия формирования отчёта в формате DOCX."""
    def make(self, data):
        filename = 'staff_{}.docx'.format(datetime.today().timestamp())
        filepath = os.path.abspath(
            os.path.join('static', 'uploads', 'reports', filename)
        )

        document = Document()

        document.add_heading('Отчёт "Сотрудники библиотеки"')

        col_names = ('Сотрудник',
                     'Занесено книг в базу данных',
                     'Обслужено читателей',
                     'Занесенные книги в базу данных',
                     'Обслуженные читатели')
        table = document.add_table(rows=1, cols=len(col_names), style="TableGrid")
        hdr_cells = table.rows[0].cells
        for idx, name in enumerate(col_names):
            paragraph = hdr_cells[idx].paragraphs[0]
            run = paragraph.add_run(name)
            run.bold = True

        for item in data:
            cells = table.add_row().cells
            cells[0].text = str(item[0])
            cells[1].text = str(item[1])
            cells[2].text = str(item[2])
            cells[3].text = str(item[3])
            cells[4].text = str(item[4])

        document.save(filepath)

        return filepath


class StaffReportPDFStrategy(StaffReportStrategy):
    """Стратегия формирования отчёта в формате PDF."""
    def make(self, data):
        filename = 'staff_{}.pdf'.format(datetime.today().timestamp())
        filepath = os.path.abspath(
            os.path.join('static', 'uploads', 'reports', filename)
        )

        response = PDFTemplateResponse(
            request=None,
            template='library_cabinet/reports/pdf_templates/staff.html',
            context={'items': data},
            cmd_options={'load-error-handling': 'ignore'})

        with open(filepath, "wb") as f:
            f.write(response.rendered_content)

        return filepath


class StaffReportODSStrategy(StaffReportStrategy):
    """Стратегия формирования отчёта в формате ODS."""
    def make(self, data):
        filename = 'orders_{}.ods'.format(datetime.today().timestamp())
        filepath = os.path.abspath(
            os.path.join('static', 'uploads', 'reports', filename)
        )

        # Формирование данных
        table_data = [['Сотрудник',
                       'Занесено книг в базу данных',
                       'Обслужено читателей',
                       'Занесенные книги в базу данных',
                       'Обслуженные читатели']]
        for item in data:
            table_data.append([
                str(item[0]),
                str(item[1]),
                str(item[2]),
                str(item[3]),
                str(item[4]),
            ])

        data = {"Sheet 1": table_data}
        save_data(filepath, data)

        return filepath


class StaffReportContext:
    """Контекст использования стратегии."""
    strategy = None
    params = {}

    def __init__(self, strategy, params):
        """Конструктор."""
        self.strategy = strategy
        self.params = params

    def get_data(self):
        """Получает данные, из которых формируется отчёт."""
        data = []

        # Формирование данных для отчёта
        for user_id in self.params.getlist('user'):
            user = User.objects.get(pk=user_id)

            # Имя сотрудника
            user_data = ['{} {}'.format(user.last_name, user.first_name)]

            # Выданные книги
            documents = DocumentModel.objects.filter(created_by_id=user_id,
                                                     created_at__gte=datetime.strptime(self.params['date_from'],
                                                                                       '%d.%m.%Y'),
                                                     created_at__lte=datetime.strptime(self.params['date_to'],
                                                                                       '%d.%m.%Y'))

            document_list = []
            for document in documents:
                document_list.append('{} ({})'.format(document, document.created_at.strftime("%d.%m.%Y %H:%M:%S")))
            user_data.append(len(document_list))

            # Действия по заказам
            order_actions = OrderAction.objects.filter(user_id=user_id,
                                                       action_date__gte=datetime.strptime(self.params['date_from'],
                                                                                          '%d.%m.%Y'),
                                                       action_date__lte=datetime.strptime(self.params['date_to'],
                                                                                          '%d.%m.%Y'))
            order_actions_list = []
            for order_action in order_actions:
                order_actions_list.append('{} - Заказ №{} ({}, {})'.format(order_action.order.user,
                                                                           order_action.order.id,
                                                                           order_action.status,
                                                                           order_action.action_date
                                                                           .strftime("%d.%m.%Y %H:%M:%S")))
            user_data.append(len(order_actions_list))

            user_data.append(", \n".join(document_list))
            user_data.append(", \n".join(order_actions_list))

            data.append(user_data)

        return data

    def get_report(self):
        """Формирует отчёт."""
        return self.strategy.make(self.get_data())
